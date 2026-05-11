from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

doc = Document()

PLOTS = Path('/home/charbel-mezeraani/cv project/evaluation/plots')

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

# ── Styles ────────────────────────────────────────────────────────────────────
styles = doc.styles

normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

NAVY = RGBColor(0x1F, 0x49, 0x7D)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)
NOTE = RGBColor(0x1A, 0x60, 0x9A)

for h_name, sz, sp_before in [
    ('Heading 1', 14, 12),
    ('Heading 2', 12,  8),
    ('Heading 3', 11,  6),
]:
    s = styles[h_name]
    s.font.name  = 'Calibri'
    s.font.size  = Pt(sz)
    s.font.bold  = True
    s.font.color.rgb = NAVY
    s.paragraph_format.space_before   = Pt(sp_before)
    s.paragraph_format.space_after    = Pt(4)
    s.paragraph_format.keep_with_next = True

# ── Helper functions ──────────────────────────────────────────────────────────
def para(text='', bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=None, color=None,
         space_after=6, space_before=0, indent=None):
    p = doc.add_paragraph(style='Normal')
    p.alignment = align
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        if size:  run.font.size      = Pt(size)
        if color: run.font.color.rgb = color
    return p

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    return p

def caption(text):
    """Centered italic figure caption."""
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(text)
    run.italic    = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY
    return p

def figure(img_path, width_inches=5.8, fig_text=''):
    """Insert image centered with caption. Falls back to placeholder if file missing."""
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    path = Path(img_path)
    if path.exists():
        run.add_picture(str(path), width=Inches(width_inches))
    else:
        run.text = f'[IMAGE NOT FOUND: {path.name}]'
        run.font.color.rgb = NOTE
        run.italic = True
    if fig_text:
        caption(fig_text)
    return p

def figure_note(text):
    """Blue italic callout for mandatory figure instructions."""
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.2)
    run = p.add_run(text)
    run.italic         = True
    run.font.size      = Pt(10)
    run.font.color.rgb = NOTE
    return p

def add_table(headers, rows, col_widths=None, font_size=10):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style     = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold           = True
        run.font.size      = Pt(font_size)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1F497D')
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        row  = table.rows[r_idx + 1]
        fill = 'EBF0F8' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(font_size)
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  fill)
            tcPr.append(shd)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def page_break():
    doc.add_page_break()

def bullet(text, level=0, size=10.5):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after       = Pt(3)
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.left_indent       = Inches(0.3 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    run = p.add_run(('•  ' if level == 0 else '–  ') + text)
    run.font.size = Pt(size)

def ref_entry(text):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after       = Pt(3)
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.left_indent       = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    run = p.add_run(text)
    run.font.size = Pt(9.5)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Real-Time Facial Emotion Recognition')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('A Progressive Deep Learning Approach: From Baseline to EfficientNetB2')
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GRAY

doc.add_paragraph()
p = doc.add_paragraph('─' * 72)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(14)

for label, value in [
    ('Course',      'Artificial Intelligence and Computer Vision'),
    ('University',  '[University Name]'),
    ('Student(s)',  '[Student Name(s)]'),
    ('Instructor',  '[Instructor Name]'),
    ('Date',        'May 2026'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{label}:  ')
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()
p = doc.add_paragraph('─' * 72)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Submitted in partial fulfillment of the requirements for the '
    'AI / Computer Vision course.\n'
    'All experiments conducted on an NVIDIA RTX 4050 Laptop GPU '
    'using TensorFlow 2.21 and Python 3.10.'
)
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GRAY

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
heading('Abstract', 1)
para(
    'Facial Emotion Recognition (FER) is a core problem in affective computing that requires '
    'automatically inferring emotional state from facial imagery. This report presents the '
    'design, implementation, evaluation, and deployment of a complete FER pipeline trained on '
    'the RAF-DB dataset across seven emotion classes: angry, disgust, fear, happy, sad, '
    'surprise, and neutral. The project follows a principled progression: a zero-shot DeepFace '
    'baseline (52.80%) establishes the task difficulty, followed by MobileNetV2 transfer '
    'learning (68.19%), a purpose-built Multi-Scale Squeeze-and-Excitation CNN (SE-CNN, '
    '69.73%), a 4-model soft-voting ensemble (70.45%), and finally EfficientNetB2 fine-tuned '
    'with a 4-phase progressive unfreezing schedule, reaching 80.02% on the 3,068-image '
    'RAF-DB test set — a +27.22% absolute improvement. Key contributions include the SE-CNN '
    'architecture combining three convolutional scales with learned channel attention; the '
    'identification of a critical training instability caused by applying photometric '
    'augmentation after per-image standardization; a systematic fine-tuning protocol with '
    'Stochastic Weight Averaging (SWA) and 8-pass Test-Time Augmentation (TTA); Grad-CAM '
    'interpretability confirming anatomically meaningful attention; and a full-stack deployment '
    'comprising a Streamlit web application, real-time webcam inference, video analysis tools, '
    'and TFLite edge export.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
heading('Table of Contents', 1)
for entry in [
    '1.  Introduction',
    '2.  Dataset and Preprocessing',
    '3.  Proposed System Architecture',
    '4.  Model Architectures',
    '5.  Training Protocol and Experimental Setup',
    '6.  Results and Analysis',
    '7.  Grad-CAM Interpretability',
    '8.  Deployment and Real-Time Application',
    '9.  Problems Faced, Solutions, and Lessons Learned',
    '10. Limitations and Future Work',
    '11. Conclusion',
    '12. References',
]:
    para(entry, space_after=4, indent=0.3)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading('1.  Introduction', 1)
heading('1.1  Problem and Motivation', 2)
para(
    'Facial expressions are among the most information-dense channels of human communication, '
    'carrying emotional content that is consistent across cultures and immediately interpretable. '
    'Automating the recognition of these expressions has direct application in human-computer '
    'interaction, healthcare monitoring, driver safety, smart classrooms, and meeting analytics. '
    'Despite these applications, automated FER is technically difficult. Primary challenges '
    'include illumination and pose variation, partial occlusion, severe class imbalance (happy '
    'and neutral dominate naturalistic data while disgust and fear are rare), low inter-class '
    'variance (fear and surprise share raised brows and widened eyes), and annotation ambiguity '
    'at low expression intensity.'
)

heading('1.2  Project Objectives', 2)
para('The four concrete objectives are:')
for item in [
    'Establish a zero-shot baseline to quantify task difficulty without training.',
    'Progress through multiple model architectures — lightweight transfer learning, a custom '
    'architecture, and a high-capacity fine-tuned model — comparing each objectively.',
    'Analyze results at the per-class level to understand what the system does and does not '
    'handle well.',
    'Deploy the system as a real-time, interpretable application.',
]:
    bullet(item)

doc.add_paragraph()
para('Table 1 — Project Contribution Summary', bold=True, space_after=4)
add_table(
    ['Contribution', 'Description', 'Why It Matters'],
    [
        ['DeepFace zero-shot baseline',
         'Evaluated directly on RAF-DB without training',
         'Establishes task difficulty; reference for all improvements'],
        ['MobileNetV2 transfer learning',
         'ImageNet pretrained backbone adapted to RAF-DB',
         'Shows value of pretrained features over random init'],
        ['Custom Multi-Scale SE-CNN',
         '3×3/5×5/7×7 branches + SE attention, trained from scratch',
         'Purpose-built for FER; demonstrates architectural reasoning'],
        ['4-model soft-voting ensemble',
         'Aggregates MobileNetV2 + SE-CNN variants',
         'Exploits model diversity to reduce prediction variance'],
        ['EfficientNetB2 fine-tuning',
         '4-phase progressive unfreezing, 80.02%',
         'Best result; shows power of large-scale pretraining'],
        ['Augmentation instability fix',
         'Removed photometric augmentation conflicting with standardization',
         'Critical engineering insight; generalizable lesson'],
        ['Grad-CAM interpretability',
         'Visualizes model attention on facial regions',
         'Validates anatomically meaningful feature learning'],
        ['Full deployment stack',
         'Streamlit, webcam, CLI, video, TFLite export',
         'Proves the system is usable beyond a training experiment'],
    ],
    col_widths=[1.6, 2.4, 2.7],
    font_size=9,
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. DATASET AND PREPROCESSING
# ══════════════════════════════════════════════════════════════════════════════
heading('2.  Dataset and Preprocessing', 1)
heading('2.1  RAF-DB Dataset', 2)
para(
    'The Real-world Affective Faces Database (RAF-DB) is collected from the internet under '
    'unconstrained real-world conditions, covering natural variation in illumination, pose, '
    'age, ethnicity, and image quality. Aligned facial images are provided at 100×100 pixels. '
    'The official test partition contains 3,068 images used for all model evaluation. Training '
    'data combines RAF-DB\'s training partition with FER2013 and a supplementary dataset. '
    'Disgust and fear are the most underrepresented classes, a fact directly reflected in '
    'their lower per-class metrics.'
)

para('Table 2 — Seven Emotion Classes in RAF-DB', bold=True, space_after=4)
add_table(
    ['Emotion', 'Primary Visual Cue', 'Training Frequency'],
    [
        ['Happy',    'Raised cheeks, visible teeth, crow\'s feet', 'High'],
        ['Neutral',  'Absence of distinctive deformations',        'High'],
        ['Sad',      'Inner brow raise, lip corner depression',     'Moderate'],
        ['Angry',    'Lowered brows, tightened lips',               'Moderate'],
        ['Surprise', 'Raised brows, widened eyes, dropped jaw',     'Moderate'],
        ['Fear',     'Raised brows, widened eyes, pulled-back lips','Low'],
        ['Disgust',  'Nose wrinkle, raised upper lip',              'Very low'],
    ],
    col_widths=[1.2, 3.5, 1.6],
)

heading('2.2  Preprocessing Pipeline', 2)
para(
    'Two separate preprocessing configurations are applied depending on the model, '
    'summarized in Table 3.'
)
para('Table 3 — Preprocessing Comparison: SE-CNN vs. EfficientNetB2', bold=True, space_after=4)
add_table(
    ['Parameter', 'SE-CNN', 'EfficientNetB2'],
    [
        ['Input size',              '48×48 px',                  '96×96 px'],
        ['Color mode',              'Grayscale (1 channel)',      'RGB (3 channels)'],
        ['Standardization',         'Per-image (zero-mean, unit-variance)', 'Per-image (zero-mean, unit-variance)'],
        ['Augmentation',            'Flip, rotate, zoom, translate', 'Flip, rotate, zoom, translate'],
        ['Photometric augmentation','Excluded',                  'Excluded'],
        ['Oversampling',            '12,000 samples per class',  'Loaded from raw RAF-DB files'],
        ['Data caching',            'NPZ cache on first run',    'No cache'],
    ],
    col_widths=[2.0, 2.2, 2.5],
)
para(
    'Grayscale for the SE-CNN reduces dimensionality and overfitting risk, as emotion is '
    'encoded in facial geometry rather than color. RGB for EfficientNetB2 is required because '
    'its pretrained ImageNet weights encode color information from the first convolution. '
    'Per-image standardization normalizes each image independently to zero mean and unit '
    'variance, making the model robust to global illumination shifts.'
)

heading('2.3  Augmentation: A Critical Engineering Finding', 2)
para(
    'Geometric augmentation (horizontal flip, ±15° rotation, ±10% zoom, ±10% translation) '
    'is applied stochastically during training. Photometric augmentations (brightness, '
    'contrast) were initially included but caused a complete training failure: accuracy '
    'remained at 14.3% — the 7-class random baseline — for the entire training run. The '
    'mechanism is a pipeline conflict: per-image standardization removes all global brightness '
    'and contrast information, so photometric augmentations applied on top introduce shifts '
    'that normalization immediately cancels, producing contradictory gradients. Removing '
    'photometric augmentation entirely resolved the collapse (see also Section 9).'
)

heading('2.4  Minority Class Oversampling', 2)
para(
    'To correct class imbalance, minority classes are replicated with augmentation until '
    'each class has 12,000 training samples (84,000 total). This ensures equal gradient '
    'contribution from all classes, critical for preventing the classifier from ignoring '
    'disgust and fear.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading('3.  Proposed System Architecture', 1)
para(
    'The system is designed as a complete computer vision pipeline. Figure 1 shows the full '
    'data flow from raw input to application output.'
)

figure(
    PLOTS / 'figure1_pipeline.png',
    width_inches=5.6,
    fig_text=(
        'Figure 1: Complete system pipeline from image/video input through face detection, '
        'preprocessing, and model inference to Grad-CAM visualization and multi-mode '
        'deployment output.'
    )
)

para(
    'Face detection uses two backends: Haar cascade (OpenCV) for fast CPU inference '
    'suitable for real-time processing, and MTCNN for higher accuracy on non-frontal or '
    'partially occluded faces. The face tracker applies temporal smoothing to prevent '
    'bounding box jitter across video frames. For offline evaluation, 8-pass TTA averages '
    'predictions over random augmented views. For real-time inference, single-pass mode is '
    'used to maintain acceptable latency. Grad-CAM is integrated directly: the gradient of '
    'the top predicted class score with respect to the final convolutional feature maps is '
    'computed, pooled to importance weights, and overlaid as a heatmap.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. MODEL ARCHITECTURES
# ══════════════════════════════════════════════════════════════════════════════
heading('4.  Model Architectures', 1)

heading('4.1  DeepFace — Zero-Shot Baseline', 2)
para(
    'DeepFace is evaluated without any training on RAF-DB, yielding 52.80% accuracy. This '
    'quantifies the cross-domain transfer gap and sets the reference point against which all '
    'subsequent improvements are measured.'
)

heading('4.2  MobileNetV2 — Lightweight Transfer Learning', 2)
para(
    'MobileNetV2 uses depthwise separable convolutions and inverted residual blocks for '
    'efficient inference. Its ImageNet-pretrained backbone is adapted by freezing the backbone '
    'and training a new classification head (Global Average Pooling → Dense → Dropout → '
    '7-class Softmax), then fine-tuning end-to-end at a reduced learning rate. SWA and '
    '8-pass TTA are applied at the final stages. Achieved: 68.19% (+15.39% over DeepFace).'
)

heading('4.3  Custom Multi-Scale SE-CNN', 2)
para(
    'The SE-CNN is designed for the spatial scale diversity inherent in facial expressions: '
    'subtle muscle tension requires fine-grained receptive fields (3×3), individual facial '
    'parts require medium fields (5×5), and holistic expression configurations require larger '
    'fields (7×7). Figure 2 shows the complete architecture.'
)

figure(
    PLOTS / 'figure2_secnn.png',
    width_inches=6.0,
    fig_text=(
        'Figure 2: Multi-Scale SE-CNN architecture. Three parallel convolutional branches '
        'capture fine, medium, and coarse expression features; their outputs are concatenated '
        'and refined by residual blocks with SE channel attention before global average pooling '
        'and classification.'
    )
)

para(
    'The Squeeze-and-Excitation (SE) block performs channel-wise attention: global average '
    'pooling (squeeze) reduces each feature map to a scalar; a two-layer MLP with sigmoid '
    'activation (excite) produces per-channel scaling weights that recalibrate feature maps '
    'by element-wise multiplication. He initialization, Batch Normalization on every '
    'convolutional layer, and residual connections complete the design. Approximately 3.2M '
    'parameters. SE-CNN v1 (best checkpoint + TTA): 69.42%; SE-CNN v2 (fine-tuned + SWA '
    '+ TTA): 69.73%.'
)

heading('4.4  EfficientNetB2 — Progressive Fine-Tuning', 2)
para(
    'EfficientNetB2 applies compound scaling to maximize accuracy per FLOP, pretrained on '
    '14 million ImageNet images. Input: 96×96 RGB. Progressive unfreezing prevents '
    'catastrophic forgetting.'
)

para('Table 4 — EfficientNetB2 Progressive Fine-Tuning Schedule', bold=True, space_after=4)
add_table(
    ['Phase', 'Layers Unfrozen', 'Epochs', 'Learning Rate', 'Val Accuracy'],
    [
        ['1 — Head only',      '0% (frozen backbone)', '15', '1×10⁻³', '57.41%'],
        ['2 — Top 40%',        '40% (deepest layers)', '20', '2×10⁻⁴', '72.72%'],
        ['3 — Full fine-tune', '100%',                 '10', '5×10⁻⁵', '74.59%'],
        ['4 — SWA refinement', '100%',                 '8',  '1×10⁻⁵', '75.16%'],
        ['Final (8-pass TTA)', '—',                    '—',  '—',       '80.02%'],
    ],
    col_widths=[1.6, 1.5, 0.75, 1.2, 1.2],
)
para(
    'Phase 1 stabilizes the classification head before modifying backbone weights. Phase 2 '
    'unfreezes the deepest layers, producing the largest single gain (+15.31%). Phase 3 '
    'extends adaptation to the full backbone. Phase 4 applies SWA to bias toward flat loss '
    'landscape regions with better generalization. The final +4.86% from TTA reflects '
    'variance reduction across 8 augmented inference passes. Mixed-precision training '
    '(float16), uint8 in-memory image storage (~300 MB RAM), and gradient clipping '
    '(clipnorm=1.0) throughout. Achieved: 80.02%.'
)

heading('4.5  Ensemble, SWA, and TTA', 2)
para(
    'The ensemble combines four models — MobileNetV2 SWA, SE-CNN v1 best, SE-CNN v2 best, '
    'and SE-CNN v2 SWA — through soft voting: 7-dimensional softmax vectors are averaged and '
    'the highest mean probability class is selected. Architectural and trajectory diversity '
    'produce partially uncorrelated errors, reducing variance. Achieved: 70.45%. SWA '
    'consistently yields +0.5–1% generalization improvement without additional data. TTA '
    'contributes +1–4.86% per model at zero training cost; single-pass mode is used for '
    'real-time inference to maintain latency.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. TRAINING PROTOCOL
# ══════════════════════════════════════════════════════════════════════════════
heading('5.  Training Protocol and Experimental Setup', 1)
para('Table 5 — Experimental Configuration Summary', bold=True, space_after=4)
add_table(
    ['Parameter', 'SE-CNN', 'EfficientNetB2'],
    [
        ['Framework',         'TensorFlow 2.21',                'TensorFlow 2.21'],
        ['Hardware',          'RTX 4050 Laptop GPU',            'RTX 4050 Laptop GPU'],
        ['Loss function',     'Categorical cross-entropy',      'Categorical cross-entropy'],
        ['Optimizer',         'Adam',                           'Adam'],
        ['LR schedule',       'Cosine decay (1×10⁻³ → ~0)',     'Manual per-phase (Table 4)'],
        ['Phase 1 epochs',    '100',                            '15'],
        ['Phase 2 epochs',    '30 (fine-tune from best ckpt)',  '20'],
        ['Phase 3 epochs',    '8 (SWA)',                        '10'],
        ['Phase 4 epochs',    '—',                              '8 (SWA)'],
        ['Early stopping',    'Patience = 15 (val loss)',       'Patience = 10 (val loss)'],
        ['Gradient clipping', '—',                              'clipnorm = 1.0'],
        ['Mixed precision',   'No',                             'float16'],
        ['Batch size',        '64',                             '32'],
        ['Evaluation set',    'Full RAF-DB test (3,068)',        'Full RAF-DB test (3,068)'],
        ['Final inference',   '8-pass TTA',                    '8-pass TTA'],
    ],
    col_widths=[2.0, 2.2, 2.5],
)
para(
    'All model checkpoints are saved on every validation accuracy improvement. Training '
    'metrics are logged per epoch to CSV files, visualized in the Streamlit Training '
    'Curves tab.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. RESULTS AND ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
heading('6.  Results and Analysis', 1)
heading('6.1  Overall Model Comparison', 2)

para(
    'Figure 3 plots the test accuracy of all six configurations on the RAF-DB test set '
    '(3,068 images, 8-pass TTA). Table 6 lists the numerical values.'
)

figure(
    PLOTS / 'model_comparison.png',
    width_inches=5.8,
    fig_text=(
        'Figure 3: Test accuracy comparison of all model configurations on the RAF-DB test set '
        '(3,068 images, 8-pass TTA). Grey: zero-shot baseline. Blue: trained models. '
        'Gold: EfficientNetB2 best model.'
    )
)

para('Table 6 — Model Accuracy Comparison', bold=True, space_after=4)
add_table(
    ['Model', 'Test Accuracy', 'Gain vs. DeepFace', 'Configuration'],
    [
        ['DeepFace pretrained (zero-shot)',       '52.80%', '—',        'Baseline'],
        ['MobileNetV2 SWA + TTA',                 '68.19%', '+15.39%',  'Transfer learning'],
        ['Custom SE-CNN v1 (best + TTA)',          '69.42%', '+16.62%',  'Trained from scratch'],
        ['Custom SE-CNN v2 (fine-tuned + TTA)',    '69.73%', '+16.93%',  'Fine-tuned from v1'],
        ['4-Model Soft-Voting Ensemble + TTA',     '70.45%', '+17.65%',  'Ensemble'],
        ['EfficientNetB2 fine-tuned + TTA',        '80.02%', '+27.22%',  'Best model'],
    ],
    col_widths=[2.4, 1.2, 1.2, 1.95],
)

heading('6.2  Ablation and Contribution Analysis', 2)
para('Table 7 — Ablation: Contribution of Each Stage', bold=True, space_after=4)
add_table(
    ['Stage', 'Accuracy', 'Marginal Gain', 'Principal Factor'],
    [
        ['DeepFace (zero-shot)',                         '52.80%', '—',        'No adaptation'],
        ['+ Task-specific fine-tuning (MobileNetV2)',    '68.19%', '+15.39%',  'ImageNet weights adapted to RAF-DB'],
        ['+ Custom architecture (SE-CNN)',                '69.73%', '+1.54%',   'Multi-scale branches + SE attention'],
        ['+ Ensemble (soft voting)',                      '70.45%', '+0.72%',   'Partially independent prediction errors'],
        ['+ High-capacity backbone (EfficientNetB2)',     '75.16%', '+4.71%',   'Richer pretrained features'],
        ['+ 8-pass TTA',                                 '80.02%', '+4.86%',   'Variance reduction at inference'],
        ['Total improvement',                            '+27.22%', '',        ''],
    ],
    col_widths=[2.5, 0.9, 1.1, 2.25],
)
para(
    'The dominant contributions are task-specific fine-tuning (+15.39%) and the combined '
    'EfficientNetB2 effect (+4.71% backbone + +4.86% TTA). The custom SE-CNN adds +1.54% '
    'over MobileNetV2, confirming genuine benefit from the purpose-built architecture. '
    'The ensemble contributes +0.72%. The central finding is that ImageNet pretraining is '
    'the single most important factor: representational richness from 14 million images '
    'cannot be replicated by architecture design alone at this dataset scale.'
)

heading('6.3  Per-Class Performance', 2)
para('Table 8 — Per-Class Precision, Recall, and F1 Score (EfficientNetB2, 80.02%)', bold=True, space_after=4)
add_table(
    ['Emotion', 'Precision', 'Recall', 'F1 Score', 'Difficulty'],
    [
        ['Happy',    '0.91', '0.92', '0.91', 'Easiest'],
        ['Surprise', '0.80', '0.76', '0.78', 'Moderate'],
        ['Angry',    '0.80', '0.64', '0.71', 'Moderate'],
        ['Sad',      '0.74', '0.81', '0.77', 'Moderate'],
        ['Neutral',  '0.72', '0.80', '0.76', 'Moderate'],
        ['Fear',     '0.63', '0.49', '0.55', 'Difficult'],
        ['Disgust',  '0.60', '0.33', '0.42', 'Hardest'],
    ],
    col_widths=[1.3, 1.1, 1.0, 1.1, 1.2],
)
para(
    'Happy achieves F1 = 0.91 because the Duchenne smile produces a large, spatially '
    'distinctive, and universally consistent facial deformation. Fear (F1 = 0.55, recall '
    '= 0.49) and surprise share the same primary muscle activations; over 50% of fear '
    'instances are misclassified, likely as surprise. Disgust (F1 = 0.42, recall = 0.33) '
    'is most problematic: the primary cue — nasal wrinkle and raised upper lip — is a '
    'subtle texture-level deformation; disgust is the rarest class in naturalistic '
    'photography; and annotation ambiguity is highest for low-intensity disgust. Angry '
    'shows a precision-recall asymmetry (P=0.80, R=0.64): the model predicts anger only '
    'when the signal is strong, misclassifying moderate anger as neutral. These results '
    'confirm that overall accuracy is an insufficient evaluation criterion — the system '
    'still misclassifies the majority of fear and disgust expressions.'
)

heading('6.4  Confusion and Failure Patterns', 2)
para('Table 9 — Representative Failure Patterns', bold=True, space_after=4)
add_table(
    ['True Label', 'Predicted Label', 'Probable Cause'],
    [
        ['Fear',    'Surprise', 'Shared brows/eyes; inner brow oblique lift not resolved'],
        ['Disgust', 'Angry',    'Low-intensity disgust with brow lowering misread as anger'],
        ['Disgust', 'Neutral',  'Very subtle nasal wrinkle below discrimination threshold'],
        ['Angry',   'Neutral',  'Moderate anger resembles concentrated neutral'],
        ['Sad',     'Neutral',  'Low-intensity sadness; inner brow raise insufficient'],
    ],
    col_widths=[1.3, 1.5, 3.95],
)
para(
    'Based on the per-class metrics, the principal confusion patterns are: Fear → Surprise '
    '(shared raised brows and widened eyes; mouth configuration is the key discriminator); '
    'Disgust → Angry or Neutral (overlapping brow-lowering muscle activation at low '
    'intensity); Angry → Neutral (low-intensity anger resembles a stern neutral face); '
    'Sad → Neutral (low-intensity sadness overlaps with neutral at the inner brow).'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. GRAD-CAM INTERPRETABILITY
# ══════════════════════════════════════════════════════════════════════════════
heading('7.  Grad-CAM Interpretability', 1)
para(
    'Gradient-weighted Class Activation Mapping (Grad-CAM) computes the gradient of the '
    'predicted class score with respect to the final convolutional feature maps, pools these '
    'gradients to per-channel importance weights, and generates a spatial heatmap '
    'highlighting the most discriminative input regions. Interpretability serves two '
    'functions in this project: as a debugging tool during development (confirming whether '
    'the model attends to facial regions rather than background), and as a transparency '
    'feature in deployment (allowing end users to verify the model\'s reasoning).'
)
para(
    'Observed attention patterns are consistent with expected facial anatomy. For happy '
    'predictions, the heatmap concentrates on the lower face — mouth, cheeks, and periorbital '
    'region — corresponding to the Duchenne smile deformation. For angry predictions, '
    'attention falls on the brow and periorbital region, consistent with brow lowering and '
    'furrowing. For difficult classes (disgust, fear), heatmaps are more diffuse, reflecting '
    'lower model certainty and the fine-grained nature of the discriminative signal. In cases '
    'of correct disgust classification, attention occasionally concentrates around the nasal '
    'region — the expected locus of disgust cues. These anatomically meaningful patterns '
    'confirm that EfficientNetB2 solves the task through genuine expression recognition '
    'rather than spurious background correlations. Grad-CAM overlays are shown in the '
    'deployment screenshots (Figure 4, Section 8).'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. DEPLOYMENT AND REAL-TIME APPLICATION
# ══════════════════════════════════════════════════════════════════════════════
heading('8.  Deployment and Real-Time Application', 1)
para(
    'The trained models are integrated into a complete deployment stack that demonstrates '
    'the system is functional beyond training experiments.'
)

para('The Streamlit web application (app/app.py) provides five interactive tabs:', space_after=4)
add_table(
    ['Tab', 'Functionality'],
    [
        ['Live Demo',
         'Image upload or webcam snapshot; per-face emotion label, confidence, probability bar chart, and Grad-CAM heatmap; optional side-by-side DeepFace comparison'],
        ['Video Analysis',
         'Frame-by-frame emotion processing; temporal timeline; stacked area chart; dominant-emotion strip'],
        ['Model Comparison',
         'Benchmark bar chart; sortable accuracy table for all six configurations'],
        ['Training Curves',
         'Accuracy and loss curves loaded live from CSV logs for all training phases'],
        ['About',
         'Architecture overview, dataset description, command reference'],
    ],
    col_widths=[1.5, 5.2],
)

# Figure 4 — deployment screenshots
fig4_path = PLOTS / 'figure4_deployment.png'
if fig4_path.exists():
    figure(
        fig4_path,
        width_inches=5.8,
        fig_text=(
            'Figure 4: Deployed Streamlit application output. Top row: happy expression detected '
            'with 100% confidence (yellow bounding box) and corresponding probability distribution. '
            'Bottom row: angry expression detected with 99.9% confidence (red bounding box).'
        )
    )
else:
    figure_note(
        '[INSERT Figure 4 here: 2×2 grid of your deployment screenshots — '
        'deploy_happy_box.png, deploy_happy_bar.png, deploy_angry_box.png, deploy_angry_bar.png. '
        'Run: python3 evaluation/plots/make_figure4_grid.py  '
        'Caption: "Figure 4: Deployed Streamlit application output showing happy detection '
        '(100% confidence, yellow box) and angry detection (99.9% confidence, red box)."]'
    )

para(
    'Real-Time Webcam (realtime/realtime_webcam.py) overlays emotion labels, probability '
    'bars, and an optional Grad-CAM heatmap on live frames. The --compare flag displays '
    'predictions from both the custom model and DeepFace side-by-side. CLI tools '
    '(detect_image.py, detect_video.py, meeting_analyzer.py) provide batch processing and '
    'meeting PDF reporting. TFLite Export with INT8 quantization produces a 4× smaller model '
    'deployable on Android smartphones and Raspberry Pi.'
)
para(
    'Deployment tradeoff: 8-pass TTA is not suitable for real-time inference (8× latency '
    'multiplier). The deployed webcam system uses single-pass inference, accepting a modest '
    'accuracy reduction for acceptable frame rates. Figure 5 shows the complete training '
    'progression for both main model families.'
)

figure(
    PLOTS / 'figure5_training_curves.png',
    width_inches=6.0,
    fig_text=(
        'Figure 5: Training and validation accuracy curves from real experiment logs. Left: '
        'MobileNetV2 across head-only and full fine-tuning phases. Right: EfficientNetB2 across '
        'all three progressive unfreezing phases with phase boundaries marked. Test accuracy '
        'with 8-pass TTA reaches 68.19% (MobileNetV2) and 80.02% (EfficientNetB2).'
    )
)

para(
    'The curves confirm expected training behavior. For MobileNetV2, validation accuracy '
    'rises steeply in Phase 1 (head training) and continues improving through full fine-tuning. '
    'For EfficientNetB2, Phase 2 (top-40% backbone unfreezing) produces the steepest accuracy '
    'gain, Phase 3 continues at a slower rate, and Phase 4 (SWA) stabilizes the model at its '
    'best generalization point. The large gap between training and validation accuracy in '
    'EfficientNetB2 is expected: training uses strong augmentation while validation does not.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. PROBLEMS FACED, SOLUTIONS, AND LESSONS LEARNED
# ══════════════════════════════════════════════════════════════════════════════
heading('9.  Problems Faced, Solutions, and Lessons Learned', 1)
para('Table 10 — Engineering Challenges, Root Causes, Solutions, and Lessons', bold=True, space_after=4)
add_table(
    ['Problem', 'Root Cause', 'Solution', 'Lesson Learned'],
    [
        ['Training collapse to 14.3%',
         'RandomBrightness/RandomContrast after per-image standardization — normalization '
         'cancels the photometric shift, contradictory gradients',
         'Remove all photometric augmentation; geometric only',
         'Preprocessing and augmentation are not independent. Validate full pipeline on '
         'small run before committing to full training'],
        ['Class imbalance',
         'Naturalistic photography skews toward happy/neutral',
         'Oversample minority classes to 12,000 samples each',
         'Oversampling is effective; focal loss is a complementary alternative'],
        ['Catastrophic forgetting',
         'High LR fine-tuning overwrites pretrained representations with noisy gradients',
         'Progressive unfreezing at decreasing learning rates',
         'Warm up task-specific layers before touching pretrained weights'],
        ['Fear/surprise confusion',
         'Both emotions share primary muscles; discriminative cues are subtle',
         'No architectural fix; partially addressed by TTA and ensemble',
         'Temporal modeling or audio fusion may be necessary'],
        ['Disgust under-recalled (R=0.33)',
         'Subtle texture cue, low training frequency, high annotation ambiguity',
         'Accepted as limitation; harder negative mining is a future option',
         'Per-class metrics should drive dataset curation decisions'],
        ['Real-time latency vs. accuracy',
         '8-pass TTA is 8× more expensive than single-pass',
         'Single-pass for real-time; TTA for offline evaluation only',
         'Benchmark and deployment accuracy should be reported separately'],
        ['EfficientNetB2 outperforms SE-CNN by ~10%',
         'SE-CNN trained on ~84K samples; EfficientNetB2 pretrained on 14M images',
         'Adopted EfficientNetB2 as primary model',
         'At academic dataset scales, a pretrained backbone almost always outperforms '
         'a custom architecture trained from scratch'],
    ],
    col_widths=[1.3, 1.65, 1.65, 2.1],
    font_size=9,
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. LIMITATIONS AND FUTURE WORK
# ══════════════════════════════════════════════════════════════════════════════
heading('10.  Limitations and Future Work', 1)
para('Current limitations:', bold=True, space_after=4)
para(
    'RAF-DB images are aligned and pre-cropped. Real webcam footage presents faces at varying '
    'resolutions, angles, and occlusion levels; the 80.02% benchmark accuracy likely '
    'overstates real-world performance. Demographic bias in web-collected training data may '
    'affect performance across age groups and ethnicities. Fear and disgust remain '
    'substantially under-performing (F1 = 0.55 and 0.42); oversampling cannot compensate '
    'for the fundamental scarcity of high-quality examples or annotation noise. Real-time '
    'deployment uses single-pass inference, sacrificing approximately 4–5% accuracy.'
)
para('Future directions:', bold=True, space_after=4)
for item in [
    'Vision Transformers (ViT / Swin): self-attention models long-range facial dependencies, '
    'potentially improving fear/surprise discrimination.',
    'Temporal modeling (LSTM / TCN): expression dynamics in video sequences are invisible '
    'in still images but strongly discriminate emotions like fear.',
    'Audio-visual fusion: vocal prosody resolves visually ambiguous expressions.',
    'Class-balanced loss functions: focal loss or class-frequency weighting as alternatives '
    'to oversampling.',
    'Larger datasets: AffectNet (1M+) or EmotioNet would provide more disgust/fear examples '
    'and greater demographic diversity.',
    'Mobile-optimized inference: knowledge distillation or neural architecture search for '
    'deployment without quantization quality loss.',
]:
    bullet(item)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
heading('11.  Conclusion', 1)
para(
    'This project designed, implemented, evaluated, and deployed a complete facial emotion '
    'recognition system progressing from a zero-shot DeepFace baseline at 52.80% to a '
    'fine-tuned EfficientNetB2 model at 80.02% — a +27.22% absolute improvement through '
    'a principled sequence of methodological choices on the RAF-DB dataset.'
)
para('The key technical findings are:', space_after=4)
for i, item in enumerate([
    'Transfer learning from ImageNet is decisive. EfficientNetB2 outperforms a 4-model '
    'ensemble of purpose-built SE-CNNs by +9.57%, confirming that 14 million pretraining '
    'images provide representational richness that architectural design alone cannot replicate '
    'at academic dataset scales.',
    'Preprocessing and augmentation interact non-trivially. Combining per-image '
    'standardization with photometric augmentation causes complete training collapse. '
    'Geometric-only augmentation is the correct choice for standardized FER pipelines.',
    'Overall accuracy is an insufficient evaluation metric. EfficientNetB2\'s 80.02% '
    'coexists with F1 = 0.42 for disgust and F1 = 0.55 for fear. Per-class analysis '
    'is essential for any FER deployment decision.',
    'Grad-CAM confirms anatomically meaningful reasoning — mouth/cheek focus for happy, '
    'brow focus for angry/surprised — validating that the learned representations are '
    'expression-grounded rather than spurious.',
    'A functional deployment is achievable on modest hardware. The complete pipeline — '
    'Streamlit web app, real-time webcam, video analysis, and TFLite edge export — '
    'runs on a laptop GPU.',
], 1):
    bullet(f'{i}. {item}')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
heading('12.  References', 1)
for ref in [
    '[1] M. Tan and Q. Le, "EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks," Proc. ICML, vol. 97, pp. 6105–6114, 2019.',
    '[2] J. Hu, L. Shen, and G. Sun, "Squeeze-and-Excitation Networks," Proc. CVPR, pp. 7132–7141, 2018.',
    '[3] R. R. Selvaraju et al., "Grad-CAM: Visual Explanations from Deep Networks via Gradient-Based Localization," Proc. ICCV, pp. 618–626, 2017.',
    '[4] M. Sandler, A. Howard, M. Zhu, A. Zhmoginov, and L.-C. Chen, "MobileNetV2: Inverted Residuals and Linear Bottlenecks," Proc. CVPR, pp. 4510–4520, 2018.',
    '[5] S. Li, W. Deng, and J. Du, "Reliable Crowdsourcing and Deep Locality-Preserving Learning for Expression Recognition in the Wild," Proc. CVPR, pp. 2852–2861, 2017. [RAF-DB]',
    '[6] S. Serengil and A. Ozpinar, "LightFace: A Hybrid Deep Face Recognition Framework," ASYU, 2020. [DeepFace]',
    '[7] K. He, X. Zhang, S. Ren, and J. Sun, "Deep Residual Learning for Image Recognition," Proc. CVPR, pp. 770–778, 2016.',
    '[8] P. Izmailov, D. Podoprikhin, T. Garipov, D. Vetrov, and A. G. Wilson, "Averaging Weights Leads to Wider Optima and Better Generalization," Proc. UAI, 2018. [SWA]',
    '[9] D. P. Kingma and J. Ba, "Adam: A Method for Stochastic Optimization," ICLR, 2015.',
    '[10] S. Ioffe and C. Szegedy, "Batch Normalization: Accelerating Deep Network Training by Reducing Internal Covariate Shift," Proc. ICML, pp. 448–456, 2015.',
    '[11] I. J. Goodfellow et al., "Challenges in Representation Learning: FER2013 Dataset," ICONIP, pp. 117–124, 2013.',
    '[12] P. Viola and M. Jones, "Rapid Object Detection Using a Boosted Cascade of Simple Features," Proc. CVPR, pp. 511–518, 2001.',
    '[13] O. Russakovsky et al., "ImageNet Large Scale Visual Recognition Challenge," IJCV, vol. 115, no. 3, pp. 211–252, 2015.',
]:
    ref_entry(ref)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = '/home/charbel-mezeraani/cv project/Emotion_Recognition_Report.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
